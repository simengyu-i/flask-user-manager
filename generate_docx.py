from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ── 辅助函数 ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p

def add_risk_table(rows):
    table = doc.add_table(rows=len(rows)+1, cols=2, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = '风险'
    hdr[1].text = '说明'
    for i, (k, v) in enumerate(rows):
        table.rows[i+1].cells[0].text = k
        table.rows[i+1].cells[1].text = v
    return table

def add_compare_table(rows, col_count=None):
    if col_count is None:
        col_count = len(rows[0]) if rows else 3
    table = doc.add_table(rows=len(rows), cols=col_count, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
    return table

def add_bold_text(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run

# ══════════════════════════════════════════
# 封面
# ══════════════════════════════════════════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Flask 用户管理平台\n安全漏洞补丁报告')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('报告日期：2026-07-07\n').font.size = Pt(11)
meta.add_run('项目地址：https://github.com/simengyu-i/flask-user-manager\n').font.size = Pt(11)
meta.add_run('漏洞总数：3 个').font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════
# 目录
# ══════════════════════════════════════════
add_heading_styled('目录', level=1)
toc = ['1. 漏洞一：硬编码凭据',
       '2. 漏洞二：字典注入与密码前端泄露',
       '3. 漏洞三：明文密码存储',
       '4. 修复总结']
for item in toc:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ══════════════════════════════════════════
# 漏洞一
# ══════════════════════════════════════════
add_heading_styled('漏洞一：硬编码凭据', level=1)

p = doc.add_paragraph()
p.add_run('严重程度：').bold = True
r = p.add_run('🔴 严重')
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
r.bold = True

add_heading_styled('漏洞说明', level=2)
doc.add_paragraph(
    '在原始代码中，用户数据（用户名、密码、邮箱、手机号、余额）以及 Flask 的 secret_key '
    '全部以字面量的形式硬编码在 app.py 文件中。任何人访问公开的 GitHub 仓库即可直接看到管理员密码和密钥。'
)

add_code('''# 原始代码 (app.py) — 硬编码问题
app.secret_key = "dev-key-2025"

USERS = {
    "admin": {
        "username": "admin",
        "password": "admin123",
        "role": "admin",
        "email": "admin@example.com",
        "phone": "13800138000",
        "balance": 99999
    },
    "alice": {
        "username": "alice",
        "password": "alice2025",
        "role": "user",
        "email": "alice@example.com",
        "phone": "13900139001",
        "balance": 100
    }
}''')

add_heading_styled('风险分析', level=2)
add_risk_table([
    ('代码泄露即数据泄露', '代码上传到公开 GitHub 后，所有人可直接看到管理员密码和联系方式'),
    ('密钥公开', 'secret_key 硬编码导致 session 可被任意伪造'),
    ('无法灵活配置', '修改端口、调试模式等需要直接改源码')
])

add_heading_styled('修复步骤', level=2)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('第 1 步：').bold = True
p.add_run('创建 config.json 和 users.json，将配置和用户数据移出源码。')

add_code('''{
    "secret_key": "dev-key-2025",
    "port": 5000,
    "debug": true
}''')

p = doc.add_paragraph()
p.add_run('第 2 步：').bold = True
p.add_run('更新 .gitignore，禁止敏感文件提交到 Git。')

add_code('''# .gitignore
users.json
config.json''')

p = doc.add_paragraph()
p.add_run('第 3 步：').bold = True
p.add_run('创建示例模板文件（config.example.json、users.example.json），密码用占位符替代。')

p = doc.add_paragraph()
p.add_run('第 4 步：').bold = True
p.add_run('修改 app.py 从 JSON 文件读取配置和数据。')

add_code('''# 修复后 (app.py)
import json
import os

config_path = os.path.join(os.path.dirname(__file__), "config.json")
users_path = os.path.join(os.path.dirname(__file__), "users.json")

with open(config_path, encoding="utf-8") as f:
    config = json.load(f)
with open(users_path, encoding="utf-8") as f:
    USERS = json.load(f)

app.secret_key = config.get("secret_key", "dev-key-2025")
port = config.get("port", 5000)
debug = config.get("debug", True)''')

add_heading_styled('修复前后对比', level=2)
add_compare_table([
    ('对比项', '修复前', '修复后'),
    ('凭据位置', 'app.py 源码中硬编码', 'users.json + config.json（已 gitignore）'),
    ('Git 泄露风险', '提交即泄露', '配置文件被 .gitignore 排除'),
    ('修改配置', '需修改源码', '直接改 JSON 文件'),
    ('示例文档', '无', 'users.example.json + config.example.json'),
])

doc.add_page_break()

# ══════════════════════════════════════════
# 漏洞二
# ══════════════════════════════════════════
add_heading_styled('漏洞二：字典注入与密码前端泄露', level=1)

p = doc.add_paragraph()
p.add_run('严重程度：').bold = True
r = p.add_run('🔴 严重')
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
r.bold = True

add_heading_styled('漏洞说明', level=2)
doc.add_paragraph(
    '原始代码中，用户登录成功后直接将完整的 USERS 字典（包含密码字段）传递给模板引擎，'
    '同时在模板中直接渲染密码字段。此外，登录表单没有任何输入校验，攻击者可以提交包含特殊字符的用户名。'
)

add_code('''# 原始代码 (app.py) — 整个字典传递给模板
user = USERS.get(username)
return render_template("index.html", user=user)

<!-- 原始代码 (templates/index.html) — 密码直接显示 -->
<tr><td>密码</td><td>{{ user.password }}</td></tr>''')

add_heading_styled('风险分析', level=2)
add_risk_table([
    ('密码显示在页面', '用户登录后自己的密码直接展示在浏览器中，被截屏即泄露'),
    ('字典数据过度暴露', '密码字段跟随字典传递到所有消费方，任何模板改动都可能意外泄露'),
    ('无输入校验', '特殊字符、原型链注入等攻击载荷可传入应用'),
])

add_heading_styled('修复步骤', level=2)

p = doc.add_paragraph()
p.add_run('第 1 步：').bold = True
p.add_run('创建 get_safe_user() 安全过滤函数，在返回用户数据给模板之前过滤掉密码字段。')

add_code('''def get_safe_user(username):
    """返回不包含密码的用户信息"""
    user = USERS.get(username)
    if user:
        return {k: v for k, v in user.items() if k != "password"}
    return None''')

doc.add_paragraph('原理分析：user.items() 返回字典的所有键值对，列表推导式生成一个新的字典，排除了 password 键。原始 USERS 字典中的数据不受影响，密码仅用于登录比对。')

p = doc.add_paragraph()
p.add_run('第 2 步：').bold = True
p.add_run('修改首页路由和登录路由，使用 safe_user 替代完整字典。')

add_code('''# 首页路由
user = get_safe_user(username) if username else None
return render_template("index.html", user=user)

# 登录路由
if user and user["password"] == password:
    session["username"] = username
    safe_user = get_safe_user(username)
    return render_template("index.html", user=safe_user)''')

p = doc.add_paragraph()
p.add_run('第 3 步：').bold = True
p.add_run('添加输入校验函数，拦截恶意输入。')

add_code('''def validate_input(value):
    """基础输入校验：只允许字母、数字、中文、@、.、-、_"""
    import re
    if not value or not isinstance(value, str):
        return False
    return bool(re.match(r'^[a-zA-Z0-9一-鿿@.\\-_]+$', value))''')

p = doc.add_paragraph()
p.add_run('第 4 步：').bold = True
p.add_run('从模板中移除密码行。')

add_code('''<!-- 修复前 -->
<tr><td>密码</td><td>{{ user.password }}</td></tr>

<!-- 修复后：整行删除，不再显示密码 -->''')

add_heading_styled('修复前后对比', level=2)
add_compare_table([
    ('对比项', '修复前', '修复后'),
    ('传递给模板的数据', '完整字典（含密码）', '安全字典（不含密码）'),
    ('页面显示的字段', '用户名、密码、邮箱、手机、角色、余额', '用户名、邮箱、手机、角色、余额'),
    ('输入校验', '无', '正则白名单校验'),
    ('注入拦截', '❌ <script> 等可传入', '✅ 非法字符返回错误提示'),
])

doc.add_page_break()

# ══════════════════════════════════════════
# 漏洞三
# ══════════════════════════════════════════
add_heading_styled('漏洞三：明文密码存储', level=1)

p = doc.add_paragraph()
p.add_run('严重程度：').bold = True
r = p.add_run('🔴 严重')
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
r.bold = True

add_heading_styled('漏洞说明', level=2)
doc.add_paragraph(
    '原始代码使用明文存储密码，并在登录时直接进行 == 字符串比对。'
    '一旦 users.json 被读取，所有密码直接暴露。'
)

add_code('''# 原始代码 — 明文存储 + == 比对
USERS = {
    "admin": {
        "password": "admin123"   # ← 明文
    }
}
if user and user["password"] == password:''')

add_heading_styled('风险分析', level=2)
add_risk_table([
    ('拖库即泄露', '如果 users.json 被读取，所有密码直接暴露'),
    ('密码复用风险', '用户往往在多个平台使用相同密码，一个泄露导致连锁沦陷'),
    ('时序攻击风险', 'Python 字符串 == 比对不是恒定时间，理论上存在时序攻击'),
])

add_heading_styled('修复步骤', level=2)

p = doc.add_paragraph()
p.add_run('第 1 步：').bold = True
p.add_run('使用 werkzeug.security 的 scrypt 哈希算法生成密码哈希值。')

add_code('''from werkzeug.security import generate_password_hash

hash_admin = generate_password_hash("admin123")
# 结果: scrypt:32768:8:1$5nBBXBovcLjbgoCR$718b1...

hash_alice = generate_password_hash("alice2025")
# 结果: scrypt:32768:8:1$VWmh3N3wSbyPFVO9$9b13b...''')

doc.add_paragraph('哈希结构说明：')
add_code('''scrypt:32768:8:1$5nBBXBovcLjbgoCR$718b1dfbba49bee9352...
├─────┴──┴─┴┴────┴──────────┴──────────────────────────┴
算法     N r p     salt (16字节随机)          hash (64字节)''')

p = doc.add_paragraph()
p.add_run('第 2 步：').bold = True
p.add_run('更新 users.json，用哈希值替换明文。')

add_code('''{
    "admin": {
        "username": "admin",
        "password": "scrypt:32768:8:1$5nBBXBovcLjbgoCR$718b1dfbba49bee93521964ffe41f01622609637e48f59219bb19f4dd3c1b4dd2cbc3c62bb3a960c76fa658d9b81055129350ae979c1ba985adb7aed0a24dd20",
        ...
    }
}''')

p = doc.add_paragraph()
p.add_run('第 3 步：').bold = True
p.add_run('修改 app.py 导入并使用 check_password_hash 替代 == 比对。')

add_code('''# 新增导入
from werkzeug.security import check_password_hash

# 修复前
if user and user["password"] == password:

# 修复后
if user and check_password_hash(user["password"], password):''')

p = doc.add_paragraph()
p.add_run('验证流程说明：').bold = True
doc.add_paragraph(
    '1. 从数据库取出哈希值，解析算法、参数、salt\n'
    '2. 用同样的参数和 salt 对用户输入的密码重新计算哈希\n'
    '3. 比较新哈希值和存储的哈希值是否一致\n'
    '4. 一致 → 登录通过；不一致 → 拒绝'
)

add_heading_styled('修复前后对比', level=2)
add_compare_table([
    ('对比项', '修复前', '修复后'),
    ('存储格式', '"admin123"（8 字符明文）', '"scrypt:..."（162 字符哈希）'),
    ('单向性', '可逆，直接可读', '不可逆，无法还原密码'),
    ('彩虹表防护', '无', '随机 salt，每用户唯一'),
    ('比对方式', '== 字符串比对', 'check_password_hash() 安全比对'),
    ('拖库后果', '密码全部暴露', '只能暴力破解，scrypt 成本极高'),
])

doc.add_page_break()

# ══════════════════════════════════════════
# 修复总结
# ══════════════════════════════════════════
add_heading_styled('修复总结', level=1)

add_heading_styled('代码变更统计', level=2)
add_compare_table([
    ('漏洞', '修改文件', '核心改动', '代码行数变化'),
    ('硬编码凭据', 'app.py、.gitignore\n新建 JSON 文件', '将硬编码数据移到外部 JSON 文件', '+45 / -22'),
    ('字典注入+密码泄露', 'app.py、index.html', '新增 get_safe_user() 过滤\n新增 validate_input() 校验', '+38 / -5'),
    ('明文密码', 'app.py、users.json', '密码改用 scrypt 哈希\n比对改用 check_password_hash()', '+2 / -1'),
])

add_heading_styled('当前剩余待修复漏洞（教学保留）', level=2)
vuln_data = [
    ('Session 伪造', 'secret_key 仍为弱密钥', '🔴 严重'),
    ('Debug 模式 RCE', '/console 可访问', '🔴 严重'),
    ('HTML 注释泄露账号', 'login.html 含管理员账号', '🔴 严重'),
    ('无暴力破解防护', '登录无频率限制', '🟠 高危'),
    ('无 CSRF 防护', '表单无 Token 校验', '🟠 高危'),
    ('无 HTTPS', '密码明文传输', '🟠 高危'),
    ('Session 永不过期', '无超时机制', '🟡 中危'),
]

table = doc.add_table(rows=len(vuln_data)+1, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = '漏洞'
hdr[1].text = '说明'
hdr[2].text = '等级'
for i, (a, b, c) in enumerate(vuln_data, start=1):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b
    table.rows[i].cells[2].text = c

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 报告结束 —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# ── 保存 ──
output_path = '/opt/flask_user_mgr/PATCH_REPORT.docx'
doc.save(output_path)
print(f'✅ Word 文档已生成：{output_path}')
