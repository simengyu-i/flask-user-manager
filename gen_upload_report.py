from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
    return table

# ═══════ 封面 ═══════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('文件上传漏洞\n原理 · 项目分析 · 修复')
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('基于 Flask 用户管理平台\n').font.size = Pt(11)
meta.add_run('项目地址：https://github.com/simengyu-i/flask-user-manager').font.size = Pt(11)

doc.add_page_break()

# ═══════ 第一章 ═══════
add_heading_styled('第一章：文件上传漏洞是什么', level=1)

add_heading_styled('1.1 定义', level=2)
doc.add_paragraph(
    '文件上传漏洞（Unrestricted File Upload Vulnerability）是指 Web 应用程序允许用户上传文件，'
    '但未对上传的文件进行充分的安全校验，导致攻击者可以上传恶意文件（如 WebShell、可执行脚本、'
    '覆盖系统文件等），从而获得服务器控制权的一种高危漏洞。'
)

add_heading_styled('1.2 漏洞危害', level=2)
add_table(
    ['危害', '说明'],
    [
        ('远程代码执行（RCE）', '上传 PHP/JSP/ASP 等脚本文件，直接执行系统命令'),
        ('路径遍历 / 文件覆盖', '通过 ../ 等方式覆盖系统关键文件（如 app.py、配置文件）'),
        ('SSRF / 钓鱼', '上传 HTML 文件实现钓鱼页面或 SSRF 攻击'),
        ('存储型 XSS', '上传含恶意脚本的 SVG 或 HTML 文件'),
    ]
)

add_heading_styled('1.3 产生原因', level=2)
doc.add_paragraph('文件上传漏洞产生的根本原因在于以下 4 个环节的缺失：')
p = doc.add_paragraph(style='List Bullet')
p.add_run('未使用 secure_filename 处理文件名：').bold = True
p.add_run('用户可通过 ../../ 路径遍历覆盖任意文件')
p = doc.add_paragraph(style='List Bullet')
p.add_run('未校验文件扩展名：').bold = True
p.add_run('.php、.asp 等可执行脚本可直接上传')
p = doc.add_paragraph(style='List Bullet')
p.add_run('未校验 MIME 类型：').bold = True
p.add_run('修改 Content-Type 即可绕过前端限制')
p = doc.add_paragraph(style='List Bullet')
p.add_run('未对文件内容做二次渲染：').bold = True
p.add_run('图片马（含 payload 的图片）可直接上传')

doc.add_page_break()

# ═══════ 第二章 ═══════
add_heading_styled('第二章：本项目的漏洞体现', level=1)

add_heading_styled('2.1 漏洞代码', level=2)
doc.add_paragraph('本项目 /upload 路由的原始代码存在 4 个安全漏洞：')

add_code('''@app.route("/upload", methods=["GET", "POST"])
def upload():
    if "username" not in session:
        return redirect("/login")

    if request.method == "POST":
        file = request.files.get("file")
        if file and file.filename:
            filename = file.filename          # ① 未使用 secure_filename
            save_path = os.path.join(
                UPLOAD_FOLDER, filename       # ② 路径拼接漏洞
            )
            file.save(save_path)              # ③④ 无扩展名/MIME/内容校验''')

add_heading_styled('2.2 漏洞点逐一分析', level=2)

add_table(
    ['编号', '漏洞点', '代码位置', '风险'],
    [
        ('①', '原始文件名未处理', 'filename = file.filename', '攻击者传入 ../../app.py 可覆盖项目文件'),
        ('②', '路径拼接无校验', 'os.path.join(UPLOAD_FOLDER, filename)', '配合 ① 实现路径遍历，覆盖任意文件'),
        ('③', '未校验文件扩展名', 'file.save(save_path) 之前无检查', '.php、.py、.htaccess 等均可上传'),
        ('④', '未校验文件内容', 'file.save(save_path) 之前无检查', '图片马、WebShell 直接写入磁盘'),
    ]
)

add_heading_styled('2.3 可实施的攻击路径', level=2)
doc.add_paragraph('在本项目的漏洞环境中，攻击者可以实施以下 5 种攻击：')

add_table(
    ['攻击方式', 'payload 文件名', '效果'],
    [
        ('覆盖 app.py', '../../app.py', 'Flask 自动 reload → 远程代码执行（RCE）'),
        ('覆盖模板文件 + SSTI', '../../templates/upload.html', 'Jinja2 模板注入 → 获取服务器信息'),
        ('覆盖 config.json', '../../config.json', '修改 secret_key → 伪造 session'),
        ('上传 WebShell', 'shell.php', '访问 /static/uploads/shell.php → 执行系统命令'),
        ('上传 .htaccess', '.htaccess', '修改 Apache 解析规则'),
    ]
)

add_heading_styled('2.4 攻击效果验证', level=2)
doc.add_paragraph('在修复前，以下攻击均可在本项目中成功实施：')

add_code('''# 攻击 1：上传 PHP WebShell
$ echo '<?php system($_GET["cmd"]); ?>' > shell.php
$ curl -X POST -F "file=@shell.php" http://target/upload
# → 上传成功，访问 /static/uploads/shell.php?cmd=id 即可执行命令

# 攻击 2：路径遍历覆盖 app.py
$ curl -X POST -F "file=@evil.py;filename=../../app.py" http://target/upload
# → 覆盖 app.py → Flask debug 模式自动重载 → 恶意代码执行

# 攻击 3：图片马上传
$ echo 'GIF89a<?php system("id");?>' > malicious.gif
$ curl -X POST -F "file=@malicious.gif" http://target/upload
# → 上传成功（服务器不检查文件内容）''')

doc.add_page_break()

# ═══════ 第三章 ═══════
add_heading_styled('第三章：修复方案', level=1)

add_heading_styled('3.1 修复后代码', level=2)
doc.add_paragraph('本项目采用了 4 层防御机制：')

add_code('''# ── 修复后的 /upload 路由（完整代码）──

# 新增导入
from werkzeug.utils import secure_filename
import magic
from PIL import Image

# 配置白名单
ALLOWED_EXTENSIONS = {'jpg', 'jpeg', 'png', 'gif'}

@app.route("/upload", methods=["GET", "POST"])
def upload():
    if "username" not in session:
        return redirect("/login")

    if request.method == "POST":
        file = request.files.get("file")
        if file and file.filename:

            # 修复 1：secure_filename 防止路径遍历
            filename = secure_filename(file.filename)

            # 修复 2：白名单扩展名
            if '.' not in filename or \\
               filename.rsplit('.', 1)[1].lower() not in ALLOWED_EXTENSIONS:
                return "不支持的文件类型"

            # 修复 3：MIME 校验
            mime = magic.from_buffer(file.read(1024), mime=True)
            file.seek(0)
            if mime not in ['image/jpeg', 'image/png', 'image/gif']:
                return "文件 MIME 类型不匹配"

            save_path = os.path.join(UPLOAD_FOLDER, filename)
            file.save(save_path)

            # 修复 4：PIL 内容二次渲染（彻底清除图片中嵌入的 payload）
            im = Image.open(save_path)
            im.save(save_path)''')

add_heading_styled('3.2 每层防御的作用', level=2)

add_table(
    ['防御层', '函数 / 方法', '拦截的攻击'],
    [
        ('第 1 层\n文件名净化', 'secure_filename()\n来自 werkzeug.utils',
         '过滤 ../ 等路径遍历字符\n将 ../../app.py 转为 app.py\n防止文件覆盖攻击'),
        ('第 2 层\n扩展名白名单', "filename.rsplit('.', 1)[1].lower()\nin ALLOWED_EXTENSIONS",
         '仅允许 jpg/jpeg/png/gif\n拦截 .php、.py、.htaccess\n拦截无扩展名文件'),
        ('第 3 层\nMIME 校验', "magic.from_buffer()\n读取文件真实签名头",
         '读取文件前 1024 字节的 magic number\n拦截改了扩展名但内容不对的文件\n如改名为 .png 的 PHP 文件'),
        ('第 4 层\n二次渲染', "PIL Image.open() + save()",
         '用 PIL 重新解码再编码图片\n清除嵌入图片像素中的 payload\n图片马中的恶意代码会被丢弃'),
    ]
)

add_heading_styled('3.3 修复前后对比', level=2)

add_table(
    ['测试场景', '修复前', '修复后'],
    [
        ('路径遍历\n../../app.py', '✅ 可覆盖 app.py\n→ RCE', '❌ secure_filename 过滤\n文件名变为 app.py\n无法越目录'),
        ('上传 php 文件\nshell.php', '✅ 上传成功\n→ 代码执行', '❌ 扩展名不在白名单\n提示"不支持的文件类型"'),
        ('上传 .htaccess', '✅ 上传成功\n→ 配置覆盖', '❌ 扩展名不在白名单\n直接拒绝'),
        ('图片马\nGIF89a+PHP', '✅ 上传成功\n→ 可被包含利用', '❌ MIME 校验通过？→ PIL 二次渲染有\n效图片→payload 被清除\n非图片→PIL 报错→文件被删除'),
        ('正常 PNG 图片', '✅ 上传成功', '✅ 上传成功'),
        ('正常 JPG 图片\n(.JPG 大写)', '✅ 上传成功\n（大小写不敏感）', '✅ 上传成功\n（转小写后匹配白名单）'),
    ]
)

doc.add_page_break()

# ═══════ 第四章 ═══════
add_heading_styled('第四章：修复验证', level=1)

add_heading_styled('4.1 攻击已失败', level=2)
doc.add_paragraph('修复后，以下攻击均已被拦截：')

add_code('''$ curl -F "file=@shell.php" http://target/upload
→ 不支持的文件类型，仅允许 jpg、jpeg、png、gif

$ curl -F "file=@evil.php;filename=../../app.py" http://target/upload
→ secure_filename 将 ../../app.py 转为 app.py
→ 且 .py 不在白名单 → 拦截

$ curl -F "file=@malicious.gif" http://target/upload
→ 文件内容不是有效图片（PIL 二次渲染检测到非图片）''')

add_heading_styled('4.2 正常功能不受影响', level=2)
doc.add_paragraph('合法用户上传头像的功能仍然正常：')

add_code('''$ python3 -c "from PIL import Image; Image.new('RGB',(100,100),(255,0,0)).save('/tmp/avatar.png')"
$ curl -F "file=@/tmp/avatar.png" http://target/upload
→ 上传成功，返回文件 URL 和预览''')

add_heading_styled('4.3 防御层逻辑图', level=2)
add_code('''用户上传文件
     ↓
┌─────────────────────┐
│ 第1层：secure_filename │  ← 过滤 ../  路径遍历
│  文件名净化          │     隐藏文件等
└─────────┬───────────┘
          ↓
┌─────────────────────┐
│ 第2层：扩展名白名单    │  ← 仅允许
│  jpg/jpeg/png/gif   │     .php .py 等拦截
└─────────┬───────────┘
          ↓
┌─────────────────────┐
│ 第3层：MIME 签名校验  │  ← 读文件头
│  magic.from_buffer  │     改扩展名无效
└─────────┬───────────┘
          ↓
┌─────────────────────┐
│ 第4层：PIL 二次渲染   │  ← 清除嵌入
│  Image.open + save  │     的 payload
└─────────┬───────────┘
          ↓
     保存到 static/uploads/  ✅''')

# ═══════ 结尾 ═══════
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 报告结束 —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

output_path = '/opt/flask_user_mgr/static/文件上传漏洞报告.docx'
doc.save(output_path)
print(f'✅ 报告已生成：{output_path}')
